import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Edit docx tables that contains patterns

file_name = "condition1.docx" # file to modify
path = "Downloads\\" # change for different file location
dl = os.path.join(os.path.expanduser("~"), path)
doc = docx.Document(dl+file_name)

# can modify if necessary
start = 1 # specify which row to start with
curr_stop = 1
intstate = 0

# variables to modify
table1 = doc.tables[0] # specify what table to edit, starting with 0
column = 5 # specify which column to edit, starting with 0
column_size = 33 # actual column size
steps = 1 # how many iterations before changing num
stop = steps + 1 

while stop != column_size:
    if curr_stop == stop:
        start += steps
        stop += steps
        intstate += 1
    if intstate % 2 == 0:
        num = "0"
    else:
        num = "1" 
    for row in range(start, stop):
        cell = table1.cell(row, column)
        cell.text = num
        for p in cell.paragraphs:
            p.alignment = 1
        curr_stop += 1
    
doc.save(dl+file_name)    